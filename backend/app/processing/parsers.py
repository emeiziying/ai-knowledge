"""
Document parsers for different file formats.
Supports PDF, Word, TXT, and Markdown files.
"""
import io
import logging
from abc import ABC, abstractmethod
from typing import Dict, Any, List, Optional
from pathlib import Path

import PyPDF2
from docx import Document as DocxDocument
import markdown
from markdown.extensions import toc

logger = logging.getLogger(__name__)


class DocumentParser(ABC):
    """Abstract base class for document parsers."""
    
    @abstractmethod
    def parse(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse document content and extract text and metadata.
        
        Args:
            file_content: Raw file content as bytes
            filename: Original filename for context
            
        Returns:
            Dictionary containing:
            - content: Extracted text content
            - metadata: Document metadata (title, structure, etc.)
            - structure: Document structure information
        """
        pass
    
    @abstractmethod
    def supports_format(self, mime_type: str, filename: str) -> bool:
        """Check if this parser supports the given file format."""
        pass


class PDFParser(DocumentParser):
    """Parser for PDF documents."""
    
    def supports_format(self, mime_type: str, filename: str) -> bool:
        return (
            mime_type == "application/pdf" or 
            filename.lower().endswith('.pdf')
        )
    
    def parse(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse PDF document and extract text content."""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Extract metadata
            metadata = {
                "title": None,
                "author": None,
                "subject": None,
                "creator": None,
                "producer": None,
                "creation_date": None,
                "modification_date": None,
                "page_count": len(pdf_reader.pages)
            }
            
            # Get PDF metadata if available
            if pdf_reader.metadata:
                metadata.update({
                    "title": pdf_reader.metadata.get("/Title"),
                    "author": pdf_reader.metadata.get("/Author"),
                    "subject": pdf_reader.metadata.get("/Subject"),
                    "creator": pdf_reader.metadata.get("/Creator"),
                    "producer": pdf_reader.metadata.get("/Producer"),
                    "creation_date": pdf_reader.metadata.get("/CreationDate"),
                    "modification_date": pdf_reader.metadata.get("/ModDate")
                })
            
            # Extract text content page by page
            pages = []
            full_content = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        pages.append({
                            "page_number": page_num + 1,
                            "content": page_text.strip()
                        })
                        full_content.append(page_text.strip())
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            # Create document structure
            structure = {
                "type": "pdf",
                "pages": pages,
                "total_pages": len(pdf_reader.pages),
                "extracted_pages": len(pages)
            }
            
            return {
                "content": "\n\n".join(full_content),
                "metadata": metadata,
                "structure": structure
            }
            
        except Exception as e:
            logger.error(f"Failed to parse PDF {filename}: {e}")
            raise ValueError(f"Failed to parse PDF document: {str(e)}")


class WordParser(DocumentParser):
    """Parser for Microsoft Word documents."""
    
    def supports_format(self, mime_type: str, filename: str) -> bool:
        return (
            mime_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword"
            ] or 
            filename.lower().endswith(('.docx', '.doc'))
        )
    
    def parse(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse Word document and extract text content."""
        try:
            doc_file = io.BytesIO(file_content)
            doc = DocxDocument(doc_file)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title,
                "author": core_props.author,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "comments": core_props.comments,
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "last_modified_by": core_props.last_modified_by,
                "revision": core_props.revision,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
            # Extract content with structure
            paragraphs = []
            tables = []
            full_content = []
            
            # Process paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    para_info = {
                        "index": i,
                        "text": paragraph.text.strip(),
                        "style": paragraph.style.name if paragraph.style else None
                    }
                    paragraphs.append(para_info)
                    full_content.append(paragraph.text.strip())
            
            # Process tables
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    table_info = {
                        "index": i,
                        "rows": len(table_data),
                        "columns": len(table_data[0]) if table_data else 0,
                        "data": table_data
                    }
                    tables.append(table_info)
                    
                    # Add table content to full text
                    table_text = "\n".join(["\t".join(row) for row in table_data])
                    full_content.append(f"[Table {i+1}]\n{table_text}")
            
            # Create document structure
            structure = {
                "type": "word",
                "paragraphs": paragraphs,
                "tables": tables,
                "total_paragraphs": len(doc.paragraphs),
                "total_tables": len(doc.tables)
            }
            
            return {
                "content": "\n\n".join(full_content),
                "metadata": metadata,
                "structure": structure
            }
            
        except Exception as e:
            logger.error(f"Failed to parse Word document {filename}: {e}")
            raise ValueError(f"Failed to parse Word document: {str(e)}")


class TextParser(DocumentParser):
    """Parser for plain text files."""
    
    def supports_format(self, mime_type: str, filename: str) -> bool:
        return (
            mime_type.startswith("text/") or 
            filename.lower().endswith(('.txt', '.text'))
        )
    
    def parse(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    content = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError("Could not decode text file with any supported encoding")
            
            # Split into lines for structure analysis
            lines = content.split('\n')
            non_empty_lines = [line.strip() for line in lines if line.strip()]
            
            # Basic metadata
            metadata = {
                "encoding": encoding,
                "total_lines": len(lines),
                "non_empty_lines": len(non_empty_lines),
                "character_count": len(content),
                "word_count": len(content.split())
            }
            
            # Simple structure analysis
            structure = {
                "type": "text",
                "lines": lines,
                "total_lines": len(lines),
                "non_empty_lines": len(non_empty_lines)
            }
            
            return {
                "content": content.strip(),
                "metadata": metadata,
                "structure": structure
            }
            
        except Exception as e:
            logger.error(f"Failed to parse text file {filename}: {e}")
            raise ValueError(f"Failed to parse text file: {str(e)}")


class MarkdownParser(DocumentParser):
    """Parser for Markdown documents."""
    
    def supports_format(self, mime_type: str, filename: str) -> bool:
        return (
            mime_type in ["text/markdown", "text/x-markdown"] or 
            filename.lower().endswith(('.md', '.markdown'))
        )
    
    def parse(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse Markdown document and extract content with structure."""
        try:
            # Decode content
            content = file_content.decode('utf-8')
            
            # Parse markdown with table of contents extension
            md = markdown.Markdown(extensions=['toc', 'tables', 'fenced_code'])
            html_content = md.convert(content)
            
            # Extract headings and structure
            headings = []
            lines = content.split('\n')
            
            for i, line in enumerate(lines):
                line = line.strip()
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    title = line.lstrip('#').strip()
                    if title:
                        headings.append({
                            "level": level,
                            "title": title,
                            "line_number": i + 1
                        })
            
            # Basic metadata
            metadata = {
                "title": headings[0]["title"] if headings else None,
                "heading_count": len(headings),
                "line_count": len(lines),
                "character_count": len(content),
                "word_count": len(content.split()),
                "has_tables": '|' in content,
                "has_code_blocks": '```' in content
            }
            
            # Document structure
            structure = {
                "type": "markdown",
                "headings": headings,
                "total_lines": len(lines),
                "html_content": html_content
            }
            
            return {
                "content": content.strip(),
                "metadata": metadata,
                "structure": structure
            }
            
        except Exception as e:
            logger.error(f"Failed to parse Markdown file {filename}: {e}")
            raise ValueError(f"Failed to parse Markdown file: {str(e)}")


class DocumentParserFactory:
    """Factory class to get appropriate parser for a document."""
    
    def __init__(self):
        self.parsers = [
            PDFParser(),
            WordParser(),
            MarkdownParser(),
            TextParser()  # Keep text parser last as fallback
        ]
    
    def get_parser(self, mime_type: str, filename: str) -> Optional[DocumentParser]:
        """Get appropriate parser for the given file type."""
        for parser in self.parsers:
            if parser.supports_format(mime_type, filename):
                return parser
        return None
    
    def parse_document(self, file_content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Parse document using appropriate parser."""
        parser = self.get_parser(mime_type, filename)
        if not parser:
            raise ValueError(f"No parser available for file type: {mime_type}")
        
        return parser.parse(file_content, filename)